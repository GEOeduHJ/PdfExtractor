"""Utility to extract PDF using pdfExtract functions and save txt/docx outputs for a given PDF file.
Usage: python extract_and_save.py "C:\path\to\file.pdf"
"""
import sys
from pathlib import Path
import base64

PDF_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else None
if not PDF_PATH or not PDF_PATH.exists():
    print('Usage: python extract_and_save.py "C:\path\to\file.pdf"')
    sys.exit(1)

import pdfExtract as pe

b = PDF_PATH.read_bytes()
pages = pe.extract_pages_from_pdf_bytes(b)
# OCR fallback already in extract_pages_from_pdf_bytes in main code, but ensure
if not pages or all((p.get('text','').strip() == '' for p in pages)):
    pages = pe.ocr_pdf_bytes(b)

normalized_pages = [pe.reflow_and_segment(p.get('text','')) for p in pages]

# Save combined txt with page headers and image markers
merged_parts = []
for i, p in enumerate(pages):
    merged_parts.append(f"=== 페이지 {i+1} ===")
    if normalized_pages[i].strip():
        merged_parts.append(normalized_pages[i])
    if p.get('has_images'):
        merged_parts.append('\n[이미지 또는 표가 포함되어 있습니다]\n')
merged_text = '\n\n'.join(merged_parts)

txt_out = PDF_PATH.with_suffix('.extracted.txt')
txt_out.write_text(merged_text, encoding='utf-8')
print(f'Wrote: {txt_out}')

# Save docx if python-docx available
if pe.HAVE_DOCX:
    from docx import Document
    doc = Document()
    for para in merged_text.split('\n\n'):
        p = para.strip()
        if not p:
            continue
        if p.startswith('=== 페이지'):
            hdr = doc.add_paragraph()
            run = hdr.add_run(p)
            run.bold = True
            continue
        if '[이미지 또는 표가 포함되어 있습니다]' in p:
            doc.add_paragraph('')
            ip = doc.add_paragraph()
            ip.add_run('[이미지 또는 표가 포함되어 있습니다]').italic = True
            doc.add_paragraph('')
            continue
        doc.add_paragraph(p)
    doc_out = PDF_PATH.with_suffix('.extracted.docx')
    doc.save(str(doc_out))
    print(f'Wrote: {doc_out}')
else:
    print('python-docx not installed; skipping docx output')
